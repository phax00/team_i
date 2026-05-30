from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "job_descriptions" / "generated_drafts" / "Generated Job Descriptions"


JOB_DESCRIPTIONS = [
    {
        "file_name": "Chief Operating Officer.docx",
        "identity": {
            "Job Title": "Chief Operating Officer",
            "Department/Function": "Operations / Manufacturing Leadership",
            "Location (Site/Region/Group)": "Group / Multi-site (Europe HQ with frequent travel to production sites and sales subsidiaries)",
            "Reports To (Role)": "Chief Executive Officer",
            "Direct Reports": "Plant Director, Group Supply Chain Director, Head of Group Quality, senior operations leaders across sites",
            "Role Level": "Enterprise/Group (Executive Leadership Level)",
        },
        "summary": "Leads the Group Operations agenda across manufacturing, supply chain execution, quality coordination, and operational transformation. Shapes the operating model needed for scale, resilience, and continuous improvement while aligning plant performance, service levels, and investment priorities with the company's growth ambitions.",
        "responsibilities": [
            ("Operations Strategy & Performance", "Define and execute the Group operations strategy across plants and fulfillment activities; set the performance cadence, target metrics, and management routines that align local execution with enterprise goals."),
            ("Plant Governance", "Oversee plant leadership across sites; ensure safe, stable, and efficient manufacturing operations, clear accountability, and disciplined escalation of production, labor, quality, and continuity issues."),
            ("Supply and Service Alignment", "Coordinate with Supply Chain, Procurement, Commercial, and Finance to balance service, cost, inventory, and capacity trade-offs in support of business priorities."),
            ("Operational Excellence", "Lead continuous improvement programs focused on throughput, waste reduction, labor productivity, planning discipline, and process standardization across the network."),
            ("Quality & Compliance Coordination", "Ensure that quality, food safety, and operational compliance expectations are embedded into the operating model and consistently executed across sites."),
            ("Capex & Transformation Prioritization", "Steer plant and operations investment priorities, including automation, expansion, asset replacement, and site improvement programs, to maximize business value and resilience."),
            ("M&A and Integration", "Lead the operations workstream for acquisitions and integrations, including Day-1 readiness, manufacturing continuity, footprint alignment, and synergy execution."),
            ("Leadership & Change", "Coach senior operations leaders, drive cross-site collaboration, and reinforce a culture of ownership, pace, and fact-based decision-making."),
            ("Illustrative KPIs", "OTIF/service level, plant OEE, yield and waste, labor productivity, inventory health, quality incidents, safety metrics, capex delivery, and cost-to-serve."),
        ],
        "org_scope": {
            "Sphere of Influence": "☑ Enterprise/Group ☐ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Drives enterprise-wide operating performance across manufacturing and supply execution; materially influences service, cost base, inventory, quality, and operational resilience.",
            "Decision Autonomy": "Makes broad operational decisions within approved strategy and budget, including network priorities, performance expectations, operating routines, and transformation sequencing; escalates major structural, legal, or investment decisions for CEO/Board approval.",
        },
        "reporting": {
            "Solid Line To": "Chief Executive Officer",
            "Matrix/Functional Reporting": "Close partnership with Chief Financial Officer on performance and investment governance; partnership with HR, IT, Quality, Procurement, R&D, and Commercial on cross-functional priorities.",
            "Direct Reports (Number, Job Titles)": "Plant Director, Group Supply Chain Director, Head of Group Quality, and selected senior operations leaders",
            "Dotted-line/Influence Relationships": "Engineering leaders, Planning leads, Logistics managers, local plant leadership teams, and external operational partners.",
        },
        "contacts": "Engages daily with executive leadership, site leaders, supply chain and quality leaders, and frequently with Finance, HR, IT, Procurement, and Commercial. Communication ranges from routine performance reviews to major escalations, investment decisions, transformation planning, and integration governance.",
        "innovation_examples": [
            "Established a Group operating rhythm with a common KPI hierarchy, escalation routines, and structured cross-site reviews for service, cost, quality, and safety.",
            "Defined the target plant governance model, including local accountability, central standards, and management routines to improve execution consistency across sites.",
            "Prioritized automation and process redesign initiatives that reduced operational waste and improved throughput without compromising quality or service continuity.",
            "Created a practical operations integration playbook for acquisitions, covering Day-1 continuity, footprint harmonization, and performance stabilization.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Engineering, Operations, Supply Chain, or related field required; Master's degree in Operations, Business, or a related discipline preferred.",
            "Work Experience": "15+ years in manufacturing or operations leadership, including multi-site responsibility in FMCG, food production, or another fast-moving industrial environment; proven experience leading scale, change, and cross-functional execution.",
            "Languages": "English (fluent); additional European language(s) preferred.",
        },
        "technical": {
            "Technical Skills": [
                "Operations Strategy & Governance - Expert",
                "Multi-site Manufacturing Leadership - Expert",
                "Performance Management & KPI Steering - Advanced",
                "Continuous Improvement & Operational Excellence - Advanced",
                "Capacity, Cost, and Service Trade-off Management - Advanced",
                "Capex Prioritization & Transformation Delivery - Advanced",
                "M&A Integration in Operations - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of plant operations, network planning, service models, procurement dependencies, and the cost/service implications of operational choices in a multi-entity environment.",
        },
        "competencies": [
            ("Leadership", "Provides clear direction across functions and sites; creates operating discipline while enabling local accountability and problem ownership."),
            ("Problem-Solving", "Balances speed, cost, risk, service, and quality in complex operational trade-offs; structures decisions with both data and pragmatism."),
            ("Communication", "Communicates performance expectations and escalations clearly to executives, plant leaders, and cross-functional stakeholders."),
            ("Adaptability & Agility", "Re-prioritizes quickly in response to supply shocks, demand shifts, incidents, and new business priorities."),
            ("Collaboration", "Builds alignment across functions that naturally operate with competing constraints, especially Finance, Commercial, Quality, Procurement, and IT."),
            ("Resilience", "Maintains pace and clarity under pressure during incidents, continuity risks, and transformation periods."),
        ],
        "problem_examples": [
            "Balance conflicting priorities between service level, inventory, and plant constraints when demand or supply conditions shift unexpectedly.",
            "Sequence network improvement and automation initiatives without destabilizing live operations or delaying business-critical deliveries.",
            "Resolve cross-site operating inconsistencies while keeping enough local flexibility for regulatory, labor, and customer-specific constraints.",
            "Lead continuity and recovery decisions during major operational disruptions, quality incidents, or supplier-related shortages.",
        ],
        "environment": {
            "Work Conditions": "~70-80% office/meeting environment; ~20-30% on-site presence in plants, warehouses, and operational locations.",
            "Travel": "Frequent within Europe for plant visits, integration programs, executive reviews, and operational escalations.",
            "Risk": "Normal occupational risk with periodic exposure to production environments (PPE as required). Accountable for operational continuity, safety alignment, and major execution risks across the network.",
        },
    },
    {
        "file_name": "Plant Director.docx",
        "identity": {
            "Job Title": "Plant Director",
            "Department/Function": "Operations / Manufacturing",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "Chief Operating Officer",
            "Direct Reports": "Production Manager, Head of Engineering and Reliability, local Quality and Logistics leadership, selected support managers",
            "Role Level": "Site Leadership Level - Director",
        },
        "summary": "Leads the end-to-end performance of the plant, ensuring safe, efficient, and compliant production while coordinating people, assets, planning, quality, and support functions. Owns the local operating rhythm and translates Group priorities into disciplined plant execution.",
        "responsibilities": [
            ("Plant Performance Leadership", "Lead day-to-day and medium-term plant performance across production, quality, engineering, logistics, and people coordination; ensure stable execution against output, service, and cost targets."),
            ("Production & Resource Coordination", "Coordinate production plans, labor deployment, asset availability, and operational priorities to deliver a reliable manufacturing schedule."),
            ("Safety, Quality & Compliance", "Ensure that food safety, quality, and operational compliance requirements are embedded into daily routines and escalated quickly when risk emerges."),
            ("Operational Escalation Management", "Own the plant response to major disruptions, including equipment downtime, staffing constraints, quality incidents, and supply interruptions."),
            ("Continuous Improvement", "Drive site-level improvement initiatives in yield, throughput, waste, labor productivity, planning discipline, and standard work."),
            ("Cross-functional Execution", "Work closely with Supply Chain, Engineering, Quality, HR, Finance, and IT to ensure that plant needs are visible and translated into practical action."),
            ("Budget & Asset Stewardship", "Manage local budget areas, operating expenses, and plant asset priorities in line with approved plans and performance expectations."),
            ("People Leadership", "Lead and develop the site leadership team; establish clarity, accountability, and a strong execution culture across the plant."),
            ("Illustrative KPIs", "OEE, schedule attainment, yield, waste, labor productivity, safety metrics, quality deviations, service performance, and local cost performance."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☐ Regional ☑ Site/Local ☐ Team",
            "Extent of Impact": "Directly influences site performance, continuity, quality, safety, cost, and employee engagement; materially affects local delivery reliability and manufacturing economics.",
            "Decision Autonomy": "Makes broad site-level decisions on resource allocation, operating priorities, escalation management, and local improvement sequencing within approved budgets and Group standards; escalates major structural or capital matters.",
        },
        "reporting": {
            "Solid Line To": "Chief Operating Officer",
            "Matrix/Functional Reporting": "Close coordination with Group Supply Chain, Group Quality, HR, Finance, and IT on functional standards and cross-site alignment.",
            "Direct Reports (Number, Job Titles)": "Production Manager, Head of Engineering and Reliability, local support leaders, and selected plant functional managers",
            "Dotted-line/Influence Relationships": "Planning, Procurement, Customer Service, R&D, and Commercial teams for supply, product, and customer-related priorities.",
        },
        "contacts": "Engages continuously with local plant leaders and supervisors, regularly with Group Operations, Quality, Supply Chain, HR, and Finance, and as needed with auditors, suppliers, and external technical partners. Communication ranges from shift escalation to strategic capacity and investment discussions.",
        "innovation_examples": [
            "Introduced a structured plant review cadence that improved transparency on throughput, downtime, labor, quality, and service constraints.",
            "Led cross-functional site initiatives to reduce waste, stabilize production planning, and improve line reliability without increasing complexity.",
            "Standardized escalation pathways and decision ownership for plant incidents, reducing response time and improving accountability.",
            "Partnered with central functions to implement practical site improvements in maintenance planning, quality routines, and digital reporting.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Engineering, Manufacturing, Operations, or related field required; advanced technical or management education preferred.",
            "Work Experience": "10+ years in plant operations, manufacturing, or industrial leadership, including responsibility for cross-functional site performance in an FMCG or food production setting.",
            "Languages": "English (fluent) and Slovak required; additional languages advantageous.",
        },
        "technical": {
            "Technical Skills": [
                "Plant Operations Leadership - Expert",
                "Production Performance Management - Advanced",
                "Quality, Food Safety, and Compliance Coordination - Advanced",
                "Labor & Resource Planning - Advanced",
                "Operational Escalation & Continuity Management - Advanced",
                "Continuous Improvement - Advanced",
                "Budget & Cost Awareness - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of plant economics, labor constraints, production scheduling, quality impact, and operational trade-offs within a larger multi-site organization.",
        },
        "competencies": [
            ("Leadership", "Creates clarity and accountability across plant functions; leads with a strong sense of pace, ownership, and operational discipline."),
            ("Problem-Solving", "Diagnoses issues quickly and mobilizes the right functions to stabilize execution and address root causes."),
            ("Communication", "Communicates clearly with plant teams, central functions, and executives during both routine reviews and time-sensitive escalations."),
            ("Adaptability & Agility", "Responds effectively to dynamic production conditions, changing priorities, and unexpected disruptions."),
            ("Collaboration", "Builds strong working relationships across functions to resolve practical operating constraints quickly."),
            ("Resilience", "Stays composed under pressure in a live production environment with time-critical decisions."),
        ],
        "problem_examples": [
            "Balance production targets with downtime, staffing constraints, and quality standards during a volatile operating week.",
            "Decide when to escalate line issues, reschedule production, or absorb short-term inefficiency to protect customer commitments.",
            "Lead cross-functional recovery from a plant incident while maintaining safety, communication discipline, and output focus.",
            "Translate central operating priorities into site action without overwhelming local teams with competing initiatives.",
        ],
        "environment": {
            "Work Conditions": "~50-60% office/meeting environment; ~40-50% presence in production and logistics areas.",
            "Travel": "Periodic within the region and to Group leadership meetings; primary focus is on-site leadership.",
            "Risk": "Regular exposure to production environments (PPE as required). Responsible for managing site-level operational, safety, quality, and continuity risks.",
        },
    },
    {
        "file_name": "Head of Engineering and Reliability.docx",
        "identity": {
            "Job Title": "Head of Engineering and Reliability",
            "Department/Function": "Engineering / Reliability / Technical Operations",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "Plant Director",
            "Direct Reports": "Project and Process Engineers, maintenance planning specialists, selected technical experts",
            "Role Level": "Site Leadership Level - Head",
        },
        "summary": "Leads technical reliability, engineering governance, and operational asset improvement for the plant. Builds the maintenance and engineering roadmap needed to improve uptime, technical stability, and delivery of process change while balancing capex, operational risk, and production realities.",
        "responsibilities": [
            ("Reliability Strategy", "Define and execute the plant reliability strategy for critical assets, maintenance planning, preventive routines, and technical risk reduction."),
            ("Engineering Governance", "Own engineering standards, technical documentation expectations, change control practices, and practical decision frameworks for modifications and upgrades."),
            ("Maintenance & Asset Health", "Partner with maintenance and operations teams to improve asset care, failure analysis, spare planning, and uptime performance."),
            ("Process Improvement Delivery", "Lead or sponsor engineering projects that improve throughput, reduce downtime, stabilize quality, or eliminate recurring technical losses."),
            ("Capex Planning & Execution", "Shape the site engineering roadmap and support business cases for capex, automation, asset replacement, and process redesign."),
            ("Root Cause Analysis", "Ensure that major breakdowns, recurring technical issues, and chronic line losses are investigated rigorously and translated into sustainable action."),
            ("Cross-functional Support", "Work closely with Production, Quality, Supply Chain, IT, and Procurement on technical priorities that affect plant performance and change delivery."),
            ("Team Leadership", "Lead and coach engineers and technical specialists; reinforce practical standards for ownership, documentation, and execution quality."),
            ("Illustrative KPIs", "planned vs. unplanned downtime, MTBF/MTTR, engineering project delivery, line reliability, maintenance compliance, and technical loss reduction."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☐ Regional ☑ Site/Local ☐ Team",
            "Extent of Impact": "Directly affects plant uptime, equipment stability, change execution quality, and technical readiness for growth and improvement.",
            "Decision Autonomy": "Makes technical and engineering decisions within plant priorities, approved budgets, and safety/compliance standards; escalates major capex, structural, or risk-bearing technical changes.",
        },
        "reporting": {
            "Solid Line To": "Plant Director",
            "Matrix/Functional Reporting": "Dotted-line collaboration with Chief Operating Officer and IT on automation, technical systems, and plant modernization priorities.",
            "Direct Reports (Number, Job Titles)": "Project and Process Engineers, planning/technical specialists, and selected engineering support roles",
            "Dotted-line/Influence Relationships": "Maintenance teams, Production leaders, Quality, Procurement, external contractors, and equipment suppliers.",
        },
        "contacts": "Engages daily with Production, Maintenance, and Engineering team members, regularly with Plant Leadership, Quality, and Supply Chain, and as needed with vendors, integrators, and external technical specialists. Communication ranges from problem-solving on live issues to longer-term technical roadmap decisions.",
        "innovation_examples": [
            "Introduced a reliability prioritization model for critical assets and chronic losses, improving focus on the highest-value technical interventions.",
            "Established stronger engineering change controls and technical documentation practices to reduce repeat errors and improve handovers.",
            "Delivered process and asset upgrades that increased line stability, reduced unplanned downtime, and supported future automation readiness.",
            "Built a practical root-cause management routine linking breakdown analysis to recurring action ownership and technical standards.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Mechanical, Electrical, Industrial, or Process Engineering required; advanced technical qualifications preferred.",
            "Work Experience": "8-12 years in engineering, maintenance, or technical operations, including leadership responsibility in a plant or industrial production environment.",
            "Languages": "English (fluent) and Slovak required.",
        },
        "technical": {
            "Technical Skills": [
                "Reliability Engineering - Expert",
                "Maintenance Strategy & Asset Health - Advanced",
                "Process Improvement & Technical Problem Solving - Advanced",
                "Engineering Change Control - Advanced",
                "Capex Planning & Project Delivery - Advanced",
                "Root Cause Analysis - Advanced",
                "Cross-functional Plant Support - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of how equipment stability, engineering standards, production planning, and quality expectations interact in a high-throughput manufacturing environment.",
        },
        "competencies": [
            ("Leadership", "Sets clear technical priorities, develops engineers effectively, and reinforces practical ownership across plant support functions."),
            ("Problem-Solving", "Works methodically through complex technical issues, separating symptoms from root causes and translating findings into sustainable action."),
            ("Communication", "Explains technical risk, options, and trade-offs clearly to both specialists and non-technical stakeholders."),
            ("Adaptability & Agility", "Balances urgent plant issues with longer-term engineering improvements and capex priorities."),
            ("Collaboration", "Partners effectively with Production, Quality, IT, and Procurement where technical issues cross functional boundaries."),
            ("Resilience", "Maintains good judgment and calm under operational pressure, especially during major breakdowns or change windows."),
        ],
        "problem_examples": [
            "Decide whether to pursue a short-term repair, a redesign, or a capital replacement when asset reliability becomes a chronic business risk.",
            "Stabilize engineering priorities when the plant faces simultaneous downtime, quality pressure, and competing project deadlines.",
            "Improve engineering documentation and change control without slowing operational responsiveness beyond what the plant can absorb.",
            "Align technical solutions with practical operator use, maintenance capability, and production realities rather than idealized design assumptions.",
        ],
        "environment": {
            "Work Conditions": "~50% office/meeting environment; ~50% in production and technical areas.",
            "Travel": "Occasional for supplier visits, technical benchmarking, and Group coordination.",
            "Risk": "Frequent exposure to production and technical environments (PPE as required). Accountable for engineering-related operational risk and safe technical execution.",
        },
    },
    {
        "file_name": "Project and Process Engineer.docx",
        "identity": {
            "Job Title": "Project and Process Engineer",
            "Department/Function": "Engineering / Process Improvement",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "Head of Engineering and Reliability",
            "Direct Reports": "No direct reports; may coordinate technicians, operators, and external vendors within projects",
            "Role Level": "Professional / Specialist",
        },
        "summary": "Supports plant performance through structured process improvement, technical project delivery, and practical problem-solving on production systems. Combines engineering analysis with hands-on implementation to improve line stability, throughput, and execution quality.",
        "responsibilities": [
            ("Process Analysis", "Analyze production processes, technical losses, bottlenecks, and workflow constraints to identify practical improvement opportunities."),
            ("Project Delivery", "Lead or support engineering projects, modifications, and technical workstreams from scoping and documentation through implementation and handover."),
            ("Operational Troubleshooting", "Partner with operations and maintenance teams on recurring technical issues, process instability, and line performance gaps."),
            ("Change Implementation", "Prepare technical change documentation, coordinate relevant stakeholders, and ensure implementation is controlled, traceable, and practical for plant use."),
            ("Performance Improvement", "Support initiatives to improve yield, changeover performance, throughput, process reliability, and operator usability."),
            ("Validation & Handover", "Verify that process changes and technical interventions achieve intended results and are embedded into standards and documentation."),
            ("Cross-functional Coordination", "Work with Quality, Production, Engineering, IT, and external vendors where process changes affect multiple workflows."),
            ("Illustrative KPIs", "project milestones, process improvement impact, reduction in recurring losses, technical change cycle time, and implementation quality."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☐ Regional ☑ Site/Local ☐ Team",
            "Extent of Impact": "Improves local manufacturing performance through targeted technical projects, process changes, and better execution discipline.",
            "Decision Autonomy": "Makes recommendations and day-to-day engineering decisions within assigned projects and technical scope; escalates larger risk, budget, or plant-wide impact decisions to engineering leadership.",
        },
        "reporting": {
            "Solid Line To": "Head of Engineering and Reliability",
            "Matrix/Functional Reporting": "Close project-based collaboration with Production, Maintenance, Quality, and Industrial IT.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Operators, shift leaders, maintenance technicians, process specialists, and external technical suppliers involved in project execution.",
        },
        "contacts": "Works daily with operators, production leaders, maintenance, and engineering colleagues; coordinates regularly with Quality, IT, and vendors when technical changes affect broader workflows. Communication ranges from practical shopfloor alignment to structured project updates and technical recommendations.",
        "innovation_examples": [
            "Led targeted process changes that reduced manual rework, improved line flow, and stabilized output on critical production steps.",
            "Built clearer implementation and validation routines for technical changes to improve adoption and reduce repeat issues after go-live.",
            "Supported engineering projects that improved line performance while keeping operator usability and maintenance practicality in scope.",
            "Translated recurring plant pain points into structured improvement actions with measurable operational impact.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Process, Industrial, Mechanical, or similar engineering discipline required.",
            "Work Experience": "3-7 years in process engineering, production engineering, continuous improvement, or a related plant environment; experience in food or FMCG manufacturing preferred.",
            "Languages": "English (working proficiency) and Slovak required.",
        },
        "technical": {
            "Technical Skills": [
                "Process Optimization - Advanced",
                "Technical Project Delivery - Advanced",
                "Root Cause Analysis - Advanced",
                "Change Documentation & Validation - Advanced",
                "Manufacturing Data Interpretation - Advanced",
                "Cross-functional Problem Solving - Advanced",
                "Continuous Improvement Methods - Advanced",
            ],
            "Business Knowledge Required": "Good understanding of production flow, operator constraints, equipment dependencies, and the practical realities of implementing change in a live plant environment.",
        },
        "competencies": [
            ("Initiative", "Takes ownership of improvement opportunities and follows them through from issue identification to operational follow-up."),
            ("Problem-Solving", "Combines structured analysis with practical engineering judgment in fast-moving operational conditions."),
            ("Communication", "Explains technical ideas and implementation steps clearly to operators, supervisors, and project stakeholders."),
            ("Adaptability", "Handles shifting priorities between urgent plant support, improvement work, and project execution."),
            ("Collaboration", "Works well with both technical and operational colleagues, building credibility through practical solutions."),
            ("Discipline", "Maintains documentation quality and implementation rigor even when timelines are tight."),
        ],
        "problem_examples": [
            "Diagnose whether a recurring output loss is driven by process design, equipment behavior, operator practice, or planning interaction.",
            "Implement a technical change without creating downstream instability or confusing ownership on the line.",
            "Balance speed of delivery with the need for documentation, validation, and sustainable handover.",
            "Translate broad improvement requests into scoped actions with measurable operational outcomes.",
        ],
        "environment": {
            "Work Conditions": "~40% office/analysis environment; ~60% in production and technical areas.",
            "Travel": "Limited; primarily site-based, with occasional vendor or benchmarking visits.",
            "Risk": "Regular exposure to production areas (PPE as required). Works in environments where technical changes must be implemented safely and with minimal operational disruption.",
        },
    },
    {
        "file_name": "Application Manager and Smart Factory.docx",
        "identity": {
            "Job Title": "Application Manager and Smart Factory",
            "Department/Function": "Information Technology / Industrial Applications",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "IT Director",
            "Direct Reports": "IFS Application Specialist and selected application or support resources as assigned",
            "Role Level": "Manager / Domain Lead",
        },
        "summary": "Leads the application landscape that connects plant operations, core business processes, and smart factory initiatives. Owns the roadmap for manufacturing-related applications, coordinates cross-functional adoption, and ensures that industrial digitalization is practical, governed, and tied to business value.",
        "responsibilities": [
            ("Application Ownership", "Own the lifecycle, roadmap, and support model for selected industrial and operational applications, including interfaces to ERP and plant systems."),
            ("Smart Factory Roadmap", "Define and prioritize smart factory use cases that improve visibility, data capture, workflow reliability, and operational decision-making."),
            ("Integration Coordination", "Coordinate application and data flows across ERP, MES, shopfloor tools, reporting layers, and plant processes in partnership with IT and operations stakeholders."),
            ("Business Partnering", "Act as the main point of contact for plant and operations teams on application needs, enhancements, adoption issues, and digitalization opportunities."),
            ("Governance & Change Management", "Ensure that new applications, enhancements, and automations follow clear ownership, documentation, training, and release practices."),
            ("Vendor & Delivery Coordination", "Coordinate external implementation partners, vendors, and internal IT resources to deliver sustainable changes on time and with acceptable supportability."),
            ("Operational Support Oversight", "Oversee issue triage, prioritization, and escalation for critical industrial application issues that affect production or planning continuity."),
            ("Illustrative KPIs", "application availability, issue resolution responsiveness, enhancement throughput, user adoption, and business impact of smart factory initiatives."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☑ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly influences application reliability, digital workflow quality, and the success of industrial digitalization across plant-related use cases.",
            "Decision Autonomy": "Makes application and roadmap decisions within approved architecture, budget, and governance boundaries; escalates major investment, security, or enterprise architecture impacts.",
        },
        "reporting": {
            "Solid Line To": "IT Director",
            "Matrix/Functional Reporting": "Close partnership with Chief Operating Officer, Plant Director, Supply Chain, and Engineering leaders on industrial use cases and digital priorities.",
            "Direct Reports (Number, Job Titles)": "IFS Application Specialist and selected application support or project resources as assigned",
            "Dotted-line/Influence Relationships": "Business Intelligence, Master Data, Production Planning, Quality, external integrators, and application vendors.",
        },
        "contacts": "Engages frequently with plant leaders, engineering, supply chain, planning, and business process owners, as well as with vendors and implementation partners. Communication ranges from issue escalations and release planning to roadmap prioritization and business case discussions.",
        "innovation_examples": [
            "Built a practical smart factory roadmap focused on use cases with immediate operational value rather than disconnected technology pilots.",
            "Improved application ownership clarity and change coordination between IT and plant teams for critical operational workflows.",
            "Introduced better support, release, and documentation routines for industrial applications connected to live production processes.",
            "Enabled faster problem resolution by making system dependencies and escalation paths clearer across ERP, plant systems, and reporting layers.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Information Systems, Engineering, Industrial IT, or related field required.",
            "Work Experience": "6-10 years in industrial applications, ERP/MES integration, manufacturing IT, or digital transformation, including ownership of cross-functional systems in an operational environment.",
            "Languages": "English (fluent) and Slovak required.",
        },
        "technical": {
            "Technical Skills": [
                "Industrial Application Ownership - Expert",
                "ERP/MES/Shopfloor Integration - Advanced",
                "Smart Factory Use Case Design - Advanced",
                "Release & Change Coordination - Advanced",
                "Vendor & Stakeholder Management - Advanced",
                "Issue Prioritization & Escalation - Advanced",
                "Business Process Translation into Digital Solutions - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of how plant execution, planning, and business process workflows depend on applications, data quality, and practical user adoption.",
        },
        "competencies": [
            ("Leadership", "Creates clear ownership and momentum across cross-functional application work without relying solely on formal authority."),
            ("Problem-Solving", "Navigates ambiguity between business pain points, system constraints, and delivery capacity to identify workable solutions."),
            ("Communication", "Translates technical dependencies into language that plant and business stakeholders can act on quickly."),
            ("Adaptability & Agility", "Balances strategic roadmap work with urgent operational support needs."),
            ("Collaboration", "Works effectively across IT, Operations, Planning, Engineering, and vendor ecosystems."),
            ("Pragmatism", "Focuses digitalization on useful, supportable, and adoptable outcomes rather than technology for its own sake."),
        ],
        "problem_examples": [
            "Decide whether an operational pain point is best solved through application change, process redesign, reporting improvement, or clearer ownership.",
            "Sequence industrial digitalization work when multiple plants and stakeholders compete for the same IT and vendor capacity.",
            "Balance application standardization with real local operational differences that affect adoption or continuity.",
            "Stabilize support and release practices for systems that are both business-critical and tightly integrated into plant operations.",
        ],
        "environment": {
            "Work Conditions": "~70% office/meeting environment; ~30% plant/site presence for workshops, support, and rollout activities.",
            "Travel": "Periodic across sites for workshops, deployments, and issue resolution.",
            "Risk": "Normal occupational risk with periodic exposure to plant areas (PPE as required). Accountable for digital workflow and application continuity risks affecting operations.",
        },
    },
    {
        "file_name": "IFS Application Specialist.docx",
        "identity": {
            "Job Title": "IFS Application Specialist",
            "Department/Function": "Information Technology / ERP Applications",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "Application Manager and Smart Factory",
            "Direct Reports": "No direct reports",
            "Role Level": "Professional / Specialist",
        },
        "summary": "Supports the reliability, usability, and continuous improvement of the IFS ERP landscape. Acts as a hands-on specialist for issue resolution, process understanding, testing, and user support across ERP workflows that affect plant, planning, finance, procurement, and operational execution.",
        "responsibilities": [
            ("ERP Support & Triage", "Investigate, prioritize, and resolve ERP-related issues or route them effectively to the right internal or external owners."),
            ("Process Understanding", "Understand the business workflows behind ERP incidents and enhancements, especially where operational, planning, procurement, or data dependencies are involved."),
            ("Testing & Release Support", "Prepare test cases, validate changes, and support controlled deployment of ERP enhancements, fixes, and process adjustments."),
            ("User Enablement", "Support key users with issue clarification, practical workarounds, training reinforcement, and improved understanding of ERP process behavior."),
            ("Documentation & Knowledge Management", "Maintain clear documentation of recurring issues, process specifics, support resolutions, and functional design decisions."),
            ("Data & Master Data Troubleshooting", "Help identify and resolve issues related to configuration, process setup, or master data that impact ERP execution quality."),
            ("Continuous Improvement Support", "Contribute to the ERP improvement backlog by translating recurring business pain points into well-scoped support or enhancement proposals."),
            ("Illustrative KPIs", "incident resolution time, issue backlog quality, testing effectiveness, enhancement readiness, and user satisfaction for supported ERP processes."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☐ Regional ☑ Site/Local ☐ Team",
            "Extent of Impact": "Improves operational continuity and user effectiveness by stabilizing the ERP workflows that underpin planning, procurement, finance, and production support processes.",
            "Decision Autonomy": "Makes day-to-day support, troubleshooting, and testing decisions within assigned functional scope; escalates architectural, investment, or broader process design decisions to application leadership.",
        },
        "reporting": {
            "Solid Line To": "Application Manager and Smart Factory",
            "Matrix/Functional Reporting": "Works closely with process owners in Finance, Procurement, Supply Chain, and Operations depending on the ERP workflow in scope.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Key users, Master Data, BI, vendor support teams, and implementation partners.",
        },
        "contacts": "Engages daily with key users, application leadership, and business process owners, and regularly with external support or implementation partners. Communication ranges from issue triage and testing updates to structured discussion of process changes and support patterns.",
        "innovation_examples": [
            "Reduced repeat support effort by documenting recurring ERP issues and codifying practical resolution paths for key user communities.",
            "Improved change readiness by strengthening testing discipline and surfacing business process impacts earlier in enhancement delivery.",
            "Connected support incidents to broader process and data quality themes, helping the organization move from reactive troubleshooting to improvement planning.",
            "Helped clarify ownership between ERP users, data stewards, and application teams for issues that crossed functional boundaries.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Information Systems, Business Informatics, Supply Chain, Finance Systems, or related field preferred.",
            "Work Experience": "3-7 years in ERP support, application management, or functional consulting, ideally with IFS or a comparable enterprise system.",
            "Languages": "English (working proficiency) and Slovak required.",
        },
        "technical": {
            "Technical Skills": [
                "ERP Functional Support - Advanced",
                "Issue Triage & Resolution - Advanced",
                "Testing & Release Validation - Advanced",
                "Process Documentation - Advanced",
                "Master Data Troubleshooting - Advanced",
                "Business User Support - Advanced",
                "Functional Change Coordination - Intermediate to Advanced",
            ],
            "Business Knowledge Required": "Working understanding of ERP-supported processes in planning, procurement, finance, inventory, and operational support, including how system and process issues interact.",
        },
        "competencies": [
            ("Ownership", "Follows issues through with discipline and keeps users informed without losing sight of quality and root causes."),
            ("Problem-Solving", "Separates data issues, user behavior, and configuration or process problems when diagnosing ERP incidents."),
            ("Communication", "Explains ERP behavior clearly to users and translates user pain points into structured technical or process language."),
            ("Adaptability", "Handles a dynamic mix of urgent support issues, testing needs, and continuous improvement work."),
            ("Collaboration", "Works effectively with users, process owners, vendors, and data-focused colleagues to resolve cross-functional problems."),
            ("Attention to Detail", "Maintains testing, documentation, and follow-up discipline in areas where small mistakes can create broad business impact."),
        ],
        "problem_examples": [
            "Determine whether an incident is caused by system behavior, configuration, training gaps, master data quality, or process misuse.",
            "Support urgent user issues without normalizing workarounds that hide deeper ERP or process problems.",
            "Validate system changes thoroughly enough to reduce business risk without slowing delivery unnecessarily.",
            "Improve ERP support quality in areas where ownership is spread across multiple business and technical stakeholders.",
        ],
        "environment": {
            "Work Conditions": "~80-90% office/remote support environment; periodic site presence for workshops, testing, or user enablement.",
            "Travel": "Occasional for site support or rollout activities.",
            "Risk": "Normal occupational risk; accountable for process and application support quality in business-critical ERP workflows.",
        },
    },
    {
        "file_name": "Business Intelligence Developer.docx",
        "identity": {
            "Job Title": "Business Intelligence Developer",
            "Department/Function": "Data and Analytics / Information Technology",
            "Location (Site/Region/Group)": "Bratislava, Slovakia",
            "Reports To (Role)": "IT Director",
            "Direct Reports": "No direct reports",
            "Role Level": "Professional / Specialist",
        },
        "summary": "Designs and delivers reporting, dashboards, and data models that help business and operational teams make better decisions. Translates business questions into usable analytics outputs while balancing data quality, usability, and maintainability across multiple stakeholder groups.",
        "responsibilities": [
            ("Dashboard & Report Delivery", "Build and maintain reports, dashboards, and self-service analytics outputs for executive, functional, and operational users."),
            ("Data Modeling", "Design and refine data models that support consistent KPI logic, reusable reporting structures, and practical business interpretation."),
            ("Requirements Translation", "Work with business stakeholders to convert reporting needs into data, logic, and visualization requirements that can be delivered and sustained."),
            ("Data Quality Support", "Identify data inconsistencies, logic gaps, and reporting issues; partner with source-system owners and data stewards to improve trust in analytics outputs."),
            ("Automation & Efficiency", "Streamline recurring reporting flows, reduce manual effort, and improve the timeliness and reliability of key business reporting."),
            ("Documentation & Knowledge Sharing", "Document logic, assumptions, and dependencies in a way that supports continuity, governance, and easier maintenance."),
            ("Cross-functional Analytics Support", "Support reporting needs across Finance, Supply Chain, Commercial, HR, or Operations depending on business priorities."),
            ("Illustrative KPIs", "dashboard adoption, report accuracy, delivery lead time, automation impact, stakeholder satisfaction, and KPI consistency."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☑ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Improves decision quality across functions by delivering trusted, usable analytics assets and helping the organization work with clearer KPI logic.",
            "Decision Autonomy": "Makes day-to-day design and technical choices within approved architecture and governance practices; escalates broader platform, source-data, or prioritization conflicts where needed.",
        },
        "reporting": {
            "Solid Line To": "IT Director",
            "Matrix/Functional Reporting": "Works closely with Finance, Supply Chain, Commercial, Operations, and Master Data stakeholders depending on active reporting priorities.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Data stewards, ERP/application owners, business controllers, planning teams, and key dashboard users.",
        },
        "contacts": "Engages frequently with business stakeholders, data owners, and IT colleagues to clarify requirements, validate KPI logic, and improve reporting usability. Communication ranges from practical report reviews to more strategic conversations about source-system limitations and data governance.",
        "innovation_examples": [
            "Replaced manually prepared reports with automated dashboards that improved availability and consistency of key operational and commercial insights.",
            "Standardized KPI definitions across reports to reduce confusion and improve alignment between business and data teams.",
            "Improved trust in analytics by surfacing data quality issues early and partnering with source-system owners on sustainable fixes.",
            "Created reusable reporting components that shortened delivery time for new dashboards and reduced maintenance complexity.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Information Systems, Data Analytics, Computer Science, Economics, or related field preferred.",
            "Work Experience": "3-7 years in business intelligence, analytics engineering, reporting, or data visualization; experience supporting business users in a multi-function environment preferred.",
            "Languages": "English (fluent); Slovak or Czech preferred.",
        },
        "technical": {
            "Technical Skills": [
                "BI Development & Visualization - Advanced",
                "SQL & Data Querying - Advanced",
                "Data Modeling & KPI Logic - Advanced",
                "Reporting Automation - Advanced",
                "Stakeholder Requirement Translation - Advanced",
                "Data Quality Diagnosis - Advanced",
                "Documentation & Reporting Governance - Intermediate to Advanced",
            ],
            "Business Knowledge Required": "Good understanding of how reporting supports operational, financial, commercial, and planning decisions, including the importance of clear KPI ownership and data definitions.",
        },
        "competencies": [
            ("Problem-Solving", "Breaks down unclear business requests into data structures, logic, and visuals that users can actually work with."),
            ("Communication", "Explains metrics, assumptions, and data limitations clearly without overwhelming business users with technical detail."),
            ("Adaptability", "Handles shifting reporting priorities and stakeholder needs while preserving delivery quality."),
            ("Collaboration", "Works effectively across functions where reporting needs and data ownership are distributed."),
            ("Discipline", "Maintains logic consistency and documentation quality in areas that are easy to fragment over time."),
            ("User Focus", "Builds outputs that are not only technically correct but also genuinely usable by their intended audience."),
        ],
        "problem_examples": [
            "Resolve conflicts between how different functions define the same KPI without losing comparability or stakeholder trust.",
            "Build useful reporting when source data is incomplete, inconsistent, or owned by multiple systems and teams.",
            "Balance fast report delivery with the need for sustainable logic, documentation, and maintainability.",
            "Decide when a reporting issue reflects visualization design, data quality, or a deeper process ownership problem.",
        ],
        "environment": {
            "Work Conditions": "~90% office/analysis environment; periodic workshop presence with business teams.",
            "Travel": "Limited; occasional for stakeholder workshops or key project meetings.",
            "Risk": "Normal occupational risk; accountable for analytics quality and reporting trust in decision-critical contexts.",
        },
    },
    {
        "file_name": "Master Data Governance Lead.docx",
        "identity": {
            "Job Title": "Master Data Governance Lead",
            "Department/Function": "Data Governance / Information Technology",
            "Location (Site/Region/Group)": "Bratislava, Slovakia",
            "Reports To (Role)": "IT Director",
            "Direct Reports": "No formal direct reports; may coordinate data stewards and business owners through governance structures",
            "Role Level": "Lead / Domain Owner",
        },
        "summary": "Leads the governance of critical master data domains to improve data quality, ownership clarity, and process consistency across the organization. Connects business process reality with governance discipline so that data standards become workable, not merely documented.",
        "responsibilities": [
            ("Data Governance Framework", "Define and maintain governance principles, ownership structures, and practical control points for key master data domains."),
            ("Domain Ownership Coordination", "Coordinate business and system owners for item, customer, supplier, and related master data areas to improve accountability and decision quality."),
            ("Data Quality Improvement", "Identify recurring data quality issues, prioritize high-impact remediation themes, and support sustainable improvement actions."),
            ("Workflow & Control Design", "Design or refine governance workflows for creation, change, review, and validation of master data across relevant systems and teams."),
            ("Business Process Alignment", "Ensure that master data governance is aligned with operational, planning, procurement, finance, and commercial process needs rather than treated as a detached control exercise."),
            ("Standards & Documentation", "Own standards, definitions, and documentation that clarify how master data should be created, maintained, and used."),
            ("Cross-functional Enablement", "Support stakeholders in understanding the business impact of poor data quality and the practical value of stronger governance routines."),
            ("Illustrative KPIs", "data quality issue recurrence, governance cycle time, master data completeness, ownership clarity, and reduction of process disruption caused by poor data."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☑ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Improves the quality and reliability of enterprise data used in planning, procurement, reporting, and operational execution across functions and sites.",
            "Decision Autonomy": "Shapes governance standards and day-to-day domain decisions within approved operating principles; escalates structural system, policy, or cross-functional prioritization conflicts where needed.",
        },
        "reporting": {
            "Solid Line To": "IT Director",
            "Matrix/Functional Reporting": "Close partnership with Finance, Supply Chain, Procurement, Commercial, and application owners on domain-specific governance decisions.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports; leads governance through influence across data stewards and business owners",
            "Dotted-line/Influence Relationships": "ERP teams, BI developers, process owners, and functional leaders impacted by master data quality.",
        },
        "contacts": "Engages regularly with process owners, application teams, analysts, and data stewards. Communication ranges from clarifying standards and ownership to facilitating issue resolution and governance decisions where data quality has material business impact.",
        "innovation_examples": [
            "Introduced a clearer ownership model for critical master data domains, reducing ambiguity around who decides, creates, approves, and maintains key records.",
            "Connected data quality issues to concrete business pain points, helping functions prioritize governance actions that mattered operationally.",
            "Built governance workflows and review routines that improved consistency without creating unnecessary administrative burden.",
            "Established stronger documentation and issue traceability for master data problems that previously resurfaced across multiple teams.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Information Systems, Business Process Management, Supply Chain, Finance Systems, or related field preferred.",
            "Work Experience": "5-8 years in master data, process governance, ERP support, or data quality-focused roles in a cross-functional business environment.",
            "Languages": "English (fluent); Slovak or Czech preferred.",
        },
        "technical": {
            "Technical Skills": [
                "Master Data Governance - Expert",
                "Data Quality Management - Advanced",
                "Cross-functional Process Design - Advanced",
                "Ownership & Workflow Definition - Advanced",
                "ERP Data Structures - Advanced",
                "Documentation & Standards Management - Advanced",
                "Stakeholder Facilitation - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of how master data affects planning, procurement, reporting, customer execution, and financial accuracy across enterprise processes.",
        },
        "competencies": [
            ("Leadership", "Drives governance discipline through influence, clarity, and structured decision-making rather than formal hierarchy alone."),
            ("Problem-Solving", "Identifies when poor data reflects a tooling issue, process issue, ownership gap, or conflicting business rule."),
            ("Communication", "Makes data governance concrete and actionable for business stakeholders who do not think in data-management terminology."),
            ("Adaptability", "Balances governance rigor with the practical constraints of changing business priorities and live operations."),
            ("Collaboration", "Builds trust across functions that often experience data issues differently but depend on the same records."),
            ("Persistence", "Follows recurring problems through to systemic improvement instead of settling for repeated short-term fixes."),
        ],
        "problem_examples": [
            "Resolve conflicts between local business practices and the need for consistent enterprise master data definitions.",
            "Improve data quality in areas where no single team fully owns the upstream business process.",
            "Design controls that are strong enough to matter but light enough that users actually follow them.",
            "Translate broad complaints about bad data into specific governance actions, domain decisions, and ownership changes.",
        ],
        "environment": {
            "Work Conditions": "~90% office/analysis environment; occasional workshops with business teams across sites.",
            "Travel": "Periodic for governance workshops and process alignment meetings.",
            "Risk": "Normal occupational risk; accountable for governance quality in data domains that materially affect process reliability and reporting trust.",
        },
    },
    {
        "file_name": "Head of Group Research and Development.docx",
        "identity": {
            "Job Title": "Head of Group Research and Development",
            "Department/Function": "Research & Development / Product Innovation",
            "Location (Site/Region/Group)": "Landgraaf, Netherlands",
            "Reports To (Role)": "Chief Executive Officer",
            "Direct Reports": "R&D managers, product development specialists, and innovation resources across sites as applicable",
            "Role Level": "Enterprise/Group (Senior Leadership Level - Head)",
        },
        "summary": "Leads the Group's product innovation agenda, shaping the pipeline, standards, and delivery model for research and development. Connects commercial opportunity, technical feasibility, quality requirements, and operational readiness to bring viable innovation into the business at pace.",
        "responsibilities": [
            ("Innovation Strategy", "Define and lead the Group innovation and product development strategy in line with growth ambitions, customer needs, and brand priorities."),
            ("Portfolio Prioritization", "Own the innovation pipeline and stage-gate governance; prioritize initiatives based on strategic fit, technical feasibility, commercial value, and execution capacity."),
            ("Product Development Leadership", "Guide formulation, product improvement, and new product development activities across categories and sites."),
            ("Cross-functional Coordination", "Work closely with Commercial, Operations, Quality, Procurement, and Finance to ensure new concepts are scalable, compliant, and economically viable."),
            ("Standards & Ways of Working", "Define development standards, documentation expectations, and decision criteria that improve consistency and execution quality across the R&D function."),
            ("Launch Readiness", "Ensure that development work translates into practical manufacturing, quality, and supply readiness for successful market introduction."),
            ("External Insight & Partnerships", "Leverage supplier, market, and technical insight to strengthen innovation quality and accelerate learning in key categories."),
            ("Illustrative KPIs", "innovation pipeline health, launch success, development cycle time, reformulation impact, and portfolio value creation."),
        ],
        "org_scope": {
            "Sphere of Influence": "☑ Enterprise/Group ☐ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Drives enterprise-level innovation priorities and materially influences growth, portfolio evolution, and technical readiness across the business.",
            "Decision Autonomy": "Makes broad R&D prioritization and functional decisions within approved strategic and investment boundaries; escalates major portfolio or investment decisions where required.",
        },
        "reporting": {
            "Solid Line To": "Chief Executive Officer",
            "Matrix/Functional Reporting": "Close partnership with COO, VP Commercial, Head of Group Quality, Procurement, and Finance on innovation feasibility, governance, and launch value.",
            "Direct Reports (Number, Job Titles)": "R&D leaders, product developers, and innovation specialists across relevant categories or locations",
            "Dotted-line/Influence Relationships": "Plant teams, Quality, Supply Chain, commercial market leads, and external technical or supplier partners.",
        },
        "contacts": "Engages regularly with executive leadership, commercial and plant stakeholders, quality and procurement teams, and external partners to align innovation priorities, technical feasibility, and launch readiness. Communication ranges from portfolio reviews to technical challenge resolution and business case discussions.",
        "innovation_examples": [
            "Created a clearer stage-gate process for innovation and renovation work, improving prioritization quality and reducing low-value pipeline clutter.",
            "Strengthened collaboration between R&D, Operations, and Quality so new concepts moved more smoothly from development into scalable production.",
            "Built a more disciplined portfolio view linking technical complexity, timeline risk, and commercial value for better decision-making.",
            "Introduced reusable development standards and documentation practices that improved continuity and cross-site knowledge sharing.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Food Science, Product Development, Chemistry, Engineering, or related field required; advanced degree preferred.",
            "Work Experience": "10+ years in R&D, product development, or innovation leadership in food, FMCG, or a similar product-led manufacturing environment.",
            "Languages": "English (fluent); additional European language(s) advantageous.",
        },
        "technical": {
            "Technical Skills": [
                "Innovation Strategy & Portfolio Leadership - Expert",
                "Product Development Governance - Advanced",
                "Cross-functional Launch Readiness - Advanced",
                "Formulation & Technical Evaluation - Advanced",
                "Stage-gate Decision Making - Advanced",
                "Supplier & External Insight Integration - Advanced",
                "Documentation & Development Standards - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of how product innovation interacts with customer needs, manufacturing feasibility, quality expectations, supply implications, and portfolio economics.",
        },
        "competencies": [
            ("Leadership", "Builds a strong innovation culture while maintaining decision discipline and practical execution focus."),
            ("Problem-Solving", "Balances creativity with technical, quality, operational, and commercial constraints to make viable portfolio decisions."),
            ("Communication", "Translates product and technical complexity into language executives and cross-functional partners can act on."),
            ("Adaptability", "Responds effectively to shifting priorities, evolving consumer needs, and trade-offs between speed and rigor."),
            ("Collaboration", "Works fluidly across functions that influence innovation success from concept to launch."),
            ("Judgment", "Knows when to push ambition and when to simplify scope to protect execution quality and timing."),
        ],
        "problem_examples": [
            "Prioritize innovation opportunities when technical complexity, commercial demand, and launch capacity are not aligned.",
            "Decide whether to reformulate, redesign, pause, or accelerate a product initiative when quality, cost, or operational constraints emerge.",
            "Build stronger process discipline in R&D without reducing creativity or slowing useful experimentation too much.",
            "Translate strategic growth ambitions into an innovation portfolio the business can realistically execute.",
        ],
        "environment": {
            "Work Conditions": "~80-90% office/lab/meeting environment; periodic visits to production and commercial locations.",
            "Travel": "Frequent within Europe for plant coordination, supplier meetings, and commercial alignment.",
            "Risk": "Normal occupational risk with periodic exposure to laboratory and production environments. Accountable for technical and portfolio risk in innovation decisions.",
        },
    },
    {
        "file_name": "Head of Group Quality.docx",
        "identity": {
            "Job Title": "Head of Group Quality",
            "Department/Function": "Quality / Food Safety / Compliance",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "Chief Operating Officer",
            "Direct Reports": "Quality managers, quality systems specialists, and selected quality experts across sites",
            "Role Level": "Enterprise/Group (Senior Leadership Level - Head)",
        },
        "summary": "Leads the Group quality agenda across food safety, compliance, quality systems, and operational quality performance. Sets the standards and governance needed to protect customers, brands, and operations while enabling scalable growth across sites and product categories.",
        "responsibilities": [
            ("Quality Strategy & Governance", "Define and maintain the Group quality and food safety strategy, standards, and management routines across all relevant entities."),
            ("Quality Systems", "Own the core quality management framework, including policy expectations, audit readiness, deviation handling, CAPA discipline, and documentation standards."),
            ("Compliance & Regulatory Coordination", "Ensure that the organization remains aligned with relevant food safety, product, and regulatory requirements across markets and sites."),
            ("Operational Quality Performance", "Partner with plants and functional leaders to improve quality performance, reduce recurring deviations, and strengthen preventive discipline."),
            ("Escalation & Incident Leadership", "Lead or coordinate the response to significant quality incidents, audit findings, or compliance risks with material business impact."),
            ("Cross-functional Integration", "Work closely with R&D, Operations, Procurement, Supply Chain, and Commercial where product, process, supplier, or customer requirements affect quality outcomes."),
            ("Capability Building", "Develop quality leadership, systems understanding, and practical quality ownership across teams and sites."),
            ("Illustrative KPIs", "audit outcomes, deviation recurrence, customer complaints, quality-related waste, CAPA closure quality, and compliance readiness."),
        ],
        "org_scope": {
            "Sphere of Influence": "☑ Enterprise/Group ☐ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly affects product safety, regulatory compliance, brand protection, operational stability, and customer confidence across the enterprise.",
            "Decision Autonomy": "Makes broad quality system and governance decisions within approved company policy and regulatory obligations; escalates major legal, customer, or enterprise risk matters.",
        },
        "reporting": {
            "Solid Line To": "Chief Operating Officer",
            "Matrix/Functional Reporting": "Close partnership with Chief Executive Officer, Plant Directors, Head of Group R&D, Procurement, and Commercial where quality issues affect broader business decisions.",
            "Direct Reports (Number, Job Titles)": "Quality leaders and specialists across sites and central quality system areas",
            "Dotted-line/Influence Relationships": "Plant quality teams, laboratory resources, external auditors, certification bodies, and supplier quality contacts.",
        },
        "contacts": "Engages regularly with plant and functional leaders, R&D, Procurement, and Supply Chain, and as needed with customers, auditors, and regulatory stakeholders. Communication ranges from routine governance and performance review to serious incident response and compliance escalation.",
        "innovation_examples": [
            "Standardized Group quality management expectations and audit routines across sites to improve consistency and readiness.",
            "Strengthened CAPA and root-cause discipline for repeat deviations, reducing recurrence and improving ownership follow-through.",
            "Connected quality governance more closely with R&D and Operations to improve launch readiness and process stability.",
            "Built clearer quality escalation pathways for incidents with potential customer, regulatory, or brand impact.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Food Science, Quality Management, Chemistry, Engineering, or related field required; advanced qualifications preferred.",
            "Work Experience": "10+ years in quality, food safety, or compliance leadership in food or FMCG manufacturing, including multi-site or Group-level responsibility.",
            "Languages": "English (fluent); additional European language(s) preferred.",
        },
        "technical": {
            "Technical Skills": [
                "Quality Governance & Systems - Expert",
                "Food Safety & Regulatory Coordination - Expert",
                "Audit Readiness & CAPA Management - Advanced",
                "Deviation & Incident Leadership - Advanced",
                "Cross-functional Quality Integration - Advanced",
                "Supplier and Process Quality Understanding - Advanced",
                "Quality Capability Building - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of how quality systems, plant execution, supplier performance, innovation choices, and customer expectations interact in a regulated manufacturing environment.",
        },
        "competencies": [
            ("Leadership", "Sets a clear quality standard while building practical ownership beyond the quality function itself."),
            ("Problem-Solving", "Separates symptom management from systemic quality improvement and drives disciplined root-cause thinking."),
            ("Communication", "Handles sensitive quality topics with clarity and credibility across operational, executive, and external audiences."),
            ("Adaptability", "Balances preventive systems work with the need to respond rapidly when incidents or audit issues arise."),
            ("Collaboration", "Works effectively across R&D, Operations, Procurement, and Commercial where quality and business outcomes intersect."),
            ("Integrity", "Maintains strong judgment and standards even under time, cost, or delivery pressure."),
        ],
        "problem_examples": [
            "Decide when a deviation reflects isolated execution failure versus a broader system weakness requiring structural intervention.",
            "Balance business urgency with the need for rigorous containment, investigation, and escalation during a quality incident.",
            "Standardize quality expectations across sites with different maturity levels without creating brittle or unrealistic controls.",
            "Strengthen quality ownership in functions that do not report into Quality but materially affect compliance outcomes.",
        ],
        "environment": {
            "Work Conditions": "~70-80% office/meeting environment; ~20-30% in plant, laboratory, or audit-related settings.",
            "Travel": "Frequent within Europe for site reviews, audits, and cross-functional quality initiatives.",
            "Risk": "Normal occupational risk with periodic exposure to production environments. Accountable for product safety, compliance, and brand-protection risks.",
        },
    },
    {
        "file_name": "Group Supply Chain Director.docx",
        "identity": {
            "Job Title": "Group Supply Chain Director",
            "Department/Function": "Supply Chain / Planning / Logistics / Customer Service",
            "Location (Site/Region/Group)": "Oberwart, Austria",
            "Reports To (Role)": "Chief Operating Officer",
            "Direct Reports": "Production Planning Manager, logistics leaders, customer service leadership, and selected supply chain experts",
            "Role Level": "Enterprise/Group (Senior Leadership Level - Director)",
        },
        "summary": "Leads the Group supply chain agenda across planning, logistics, customer service, and supply coordination. Balances service, inventory, capacity, and execution constraints to create a more reliable, responsive, and scalable supply model across markets and sites.",
        "responsibilities": [
            ("Supply Chain Strategy", "Define and execute the Group supply chain strategy across planning, logistics, and customer service in line with business growth and service ambitions."),
            ("Planning Governance", "Establish the planning model, decision forums, and KPI routines needed to balance demand, production, inventory, and service trade-offs consistently."),
            ("Logistics & Fulfillment Oversight", "Own the operating standards and performance expectations for logistics, transport coordination, and order fulfillment support."),
            ("Customer Service Integration", "Ensure that customer service and supply execution work as one operating model with clear escalation pathways and realistic service commitments."),
            ("Inventory & Service Optimization", "Improve inventory health, planning discipline, and supply responsiveness while protecting cash, service, and plant stability."),
            ("Cross-functional Coordination", "Work closely with Commercial, Procurement, Operations, Finance, and IT on supply decisions, constraints, and operating improvements."),
            ("Transformation & Standardization", "Lead programs that improve supply chain visibility, planning maturity, process consistency, and digital support for decision-making."),
            ("Illustrative KPIs", "OTIF, inventory health, forecast bias/accuracy, planning adherence, logistics service, order issue rate, and supply chain cost performance."),
        ],
        "org_scope": {
            "Sphere of Influence": "☑ Enterprise/Group ☐ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly influences service performance, inventory, planning quality, logistics efficiency, and operational responsiveness across the company.",
            "Decision Autonomy": "Makes broad supply chain and planning decisions within approved budgets and enterprise priorities; escalates major structural or investment decisions where needed.",
        },
        "reporting": {
            "Solid Line To": "Chief Operating Officer",
            "Matrix/Functional Reporting": "Close partnership with Commercial, Procurement, Plant Leadership, Finance, and IT on service, inventory, and planning trade-offs.",
            "Direct Reports (Number, Job Titles)": "Production Planning Manager, logistics and customer service leaders, and selected Group supply chain experts",
            "Dotted-line/Influence Relationships": "Country commercial teams, plant planners, warehouse teams, transport partners, and supply chain analysts.",
        },
        "contacts": "Engages daily with supply chain, planning, logistics, and customer service teams, and frequently with Commercial, Operations, Procurement, and Finance. Communication ranges from routine supply alignment to business-critical service escalations and network trade-off decisions.",
        "innovation_examples": [
            "Introduced clearer Group planning forums and decision ownership to improve alignment between demand, capacity, and service priorities.",
            "Improved supply chain visibility through more consistent KPI definitions and better escalation routines for shortages, delays, and order risk.",
            "Standardized logistics and customer service workflows to reduce avoidable friction between external commitments and internal supply capability.",
            "Built a stronger foundation for planning and execution by connecting process discipline with digital reporting and clearer ownership.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Supply Chain, Operations, Logistics, Engineering, or related field required; advanced qualifications preferred.",
            "Work Experience": "10+ years in supply chain, planning, logistics, or customer service leadership, including cross-functional and multi-site responsibility in a manufacturing environment.",
            "Languages": "English (fluent); German and/or additional European language(s) advantageous.",
        },
        "technical": {
            "Technical Skills": [
                "Supply Chain Strategy & Governance - Expert",
                "Planning Leadership - Advanced",
                "Inventory & Service Trade-off Management - Advanced",
                "Logistics & Fulfillment Oversight - Advanced",
                "Customer Service Integration - Advanced",
                "Cross-functional Supply Decision Making - Advanced",
                "Supply Chain Transformation - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of how demand, planning, plant capacity, procurement constraints, customer commitments, and financial objectives interact in an integrated supply model.",
        },
        "competencies": [
            ("Leadership", "Aligns multiple supply functions around one operating model and clear service-oriented priorities."),
            ("Problem-Solving", "Makes trade-offs between service, stock, cost, and capacity with a strong sense of business context and timing."),
            ("Communication", "Communicates supply constraints, options, and decisions clearly to both operational and commercial stakeholders."),
            ("Adaptability", "Responds effectively to volatility in demand, supply, transport, and plant conditions."),
            ("Collaboration", "Builds strong connections across Commercial, Operations, Procurement, IT, and Finance where supply decisions have shared consequences."),
            ("Composure", "Maintains structure and focus during high-pressure supply disruptions or customer-impacting events."),
        ],
        "problem_examples": [
            "Balance inventory reduction goals with unstable demand, production variability, and service expectations from commercial teams.",
            "Decide how to allocate constrained product across customers, markets, or channels under limited supply.",
            "Improve planning discipline where source data, ownership, and operating routines are fragmented across functions.",
            "Stabilize customer service and logistics performance without over-customizing workflows for every local need.",
        ],
        "environment": {
            "Work Conditions": "~80% office/meeting environment; ~20% visits to plant, warehouse, or logistics locations.",
            "Travel": "Frequent within Europe for supply reviews, site alignment, and partner meetings.",
            "Risk": "Normal occupational risk with periodic exposure to operational environments. Accountable for service, inventory, and continuity risks across the supply network.",
        },
    },
    {
        "file_name": "Production Planning Manager.docx",
        "identity": {
            "Job Title": "Production Planning Manager",
            "Department/Function": "Supply Chain / Production Planning",
            "Location (Site/Region/Group)": "Banska Bystrica, Slovakia",
            "Reports To (Role)": "Group Supply Chain Director",
            "Direct Reports": "Production planners and planning support resources as assigned",
            "Role Level": "Manager / Functional Lead",
        },
        "summary": "Leads production planning for the site or assigned manufacturing scope, turning supply requirements into achievable schedules that balance demand, capacity, inventory, and operational constraints. Coordinates closely with plant, supply chain, and customer-facing functions to keep execution realistic and responsive.",
        "responsibilities": [
            ("Production Scheduling", "Own the production planning process and scheduling logic for assigned lines or site scope, translating supply requirements into executable plans."),
            ("Capacity & Constraint Management", "Balance demand priorities with line capacity, labor, maintenance windows, material availability, and operational limitations."),
            ("Cross-functional Planning Alignment", "Coordinate with plant leadership, supply chain, procurement, customer service, and commercial stakeholders on planning assumptions, risks, and trade-offs."),
            ("Inventory & Service Support", "Support decisions that balance stock position, customer commitments, manufacturing efficiency, and production stability."),
            ("Planning Discipline", "Improve planning routines, data quality, and schedule adherence to reduce last-minute instability and reactive decision-making."),
            ("Escalation Management", "Highlight plan risks early and structure escalation options when demand, supply, or operational constraints cannot all be met simultaneously."),
            ("Continuous Improvement", "Identify opportunities to improve planning accuracy, sequencing, schedule robustness, and communication quality across the process."),
            ("Illustrative KPIs", "schedule attainment, adherence to plan, inventory health, planning responsiveness, service impact, and reduction of avoidable replanning."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☐ Regional ☑ Site/Local ☐ Team",
            "Extent of Impact": "Directly influences production stability, schedule realism, inventory health, and service performance through better planning decisions and coordination.",
            "Decision Autonomy": "Makes day-to-day planning and sequencing decisions within agreed priorities and governance; escalates broader service, inventory, or capacity conflicts when trade-offs exceed local authority.",
        },
        "reporting": {
            "Solid Line To": "Group Supply Chain Director",
            "Matrix/Functional Reporting": "Works closely with Plant Director, Production, Procurement, Customer Service, and Commercial on local execution realities and plan risk.",
            "Direct Reports (Number, Job Titles)": "Production planners and planning support resources as assigned",
            "Dotted-line/Influence Relationships": "Warehouse, Logistics, Engineering, and Quality teams where production changes affect broader execution quality.",
        },
        "contacts": "Engages daily with planners, production leadership, procurement, and customer-facing teams to align schedules and manage constraints. Communication ranges from routine schedule coordination to urgent escalation on capacity, service, or material shortages.",
        "innovation_examples": [
            "Improved planning stability by introducing clearer freeze windows, prioritization rules, and earlier visibility of operational constraints.",
            "Reduced avoidable replanning through better coordination between planning, procurement, and plant execution teams.",
            "Built more practical planning routines that reflected actual line behavior and maintenance realities instead of purely theoretical scheduling logic.",
            "Strengthened planning communication so stakeholders understood both the plan and the consequences of late change requests.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Supply Chain, Operations, Industrial Engineering, Business, or related field preferred.",
            "Work Experience": "5+ years in production planning, supply planning, scheduling, or a related manufacturing coordination role.",
            "Languages": "English (working proficiency) and Slovak required.",
        },
        "technical": {
            "Technical Skills": [
                "Production Planning & Scheduling - Expert",
                "Capacity Balancing - Advanced",
                "Inventory & Service Trade-off Management - Advanced",
                "Cross-functional Coordination - Advanced",
                "Planning Data Interpretation - Advanced",
                "Escalation Structuring - Advanced",
                "Continuous Improvement in Planning Processes - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of line behavior, supply constraints, customer implications, and the practical costs of schedule instability in a production environment.",
        },
        "competencies": [
            ("Ownership", "Maintains a strong grip on plan quality and follows through when assumptions or constraints change."),
            ("Problem-Solving", "Structures planning choices clearly when not all customer, inventory, and operational goals can be met simultaneously."),
            ("Communication", "Explains schedule logic, changes, and planning risks clearly to both operational and non-operational stakeholders."),
            ("Adaptability", "Responds well to frequent changes while preserving as much schedule stability as possible."),
            ("Collaboration", "Works effectively with functions that each experience planning pressure differently."),
            ("Discipline", "Maintains planning rigor and documentation even in fast-changing circumstances."),
        ],
        "problem_examples": [
            "Re-plan effectively when material shortages, maintenance downtime, and urgent demand changes hit the same production horizon.",
            "Decide when to protect schedule stability versus when to accept disruption to meet a higher-priority business need.",
            "Improve plan accuracy in an environment where source data and late change requests are both imperfect.",
            "Translate plant realities into planning decisions that customer-facing teams can understand and support.",
        ],
        "environment": {
            "Work Conditions": "~75-85% office/analysis environment; ~15-25% presence in production or planning review settings.",
            "Travel": "Limited; primarily site-based with periodic cross-site or Group coordination meetings.",
            "Risk": "Normal occupational risk with periodic exposure to production environments. Accountable for planning quality and execution stability.",
        },
    },
    {
        "file_name": "HR Manager - Banska Bystrica Plant.docx",
        "identity": {
            "Job Title": "HR Manager",
            "Department/Function": "Human Resources",
            "Location (Site/Region/Group)": "Banska Bystrica Plant, Slovakia",
            "Reports To (Role)": "HR Director",
            "Direct Reports": "HR specialists, HR administrators, and onboarding support resources as assigned",
            "Role Level": "Site/Local Management",
        },
        "summary": "Leads site-level HR execution for the Banska Bystrica plant, ensuring that workforce planning, employee relations, onboarding, and people processes support stable plant operations. Acts as the main HR partner for local leadership while aligning site decisions with Group HR standards and legal requirements.",
        "responsibilities": [
            ("Site Workforce Planning", "Coordinate hiring, workforce planning, and staffing priorities for the plant in line with production needs, shift models, and growth requirements."),
            ("Employee Relations & Engagement", "Handle employee relations cases, support local leadership on people issues, and reinforce engagement and communication practices appropriate for a manufacturing environment."),
            ("Onboarding & Employee Lifecycle", "Ensure structured onboarding, probation follow-up, employee documentation, and local lifecycle processes are executed consistently and compliantly."),
            ("HR Process Execution", "Implement HR policies, processes, and reporting routines at the site level while identifying where local adaptation is needed to support operational realities."),
            ("Manager Support", "Coach line managers and supervisors on performance conversations, absence topics, employee handling, and practical people leadership."),
            ("Compliance & Labor Discipline", "Ensure alignment with local labor regulations, internal HR policies, and site-specific compliance requirements affecting people processes."),
            ("Change & Plant Support", "Support organizational changes, workforce adjustments, and local integration needs linked to production changes or restructuring."),
            ("Illustrative KPIs", "time-to-fill for plant roles, onboarding completion, absenteeism trends, employee relations case closure, and local HR process compliance."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☐ Regional ☑ Site/Local ☐ Team",
            "Extent of Impact": "Directly affects plant workforce stability, compliance, employee experience, and day-to-day leadership effectiveness.",
            "Decision Autonomy": "Makes local HR operational decisions within approved policy and budget boundaries; escalates structural, legal, or higher-risk people matters to HR Director and relevant site leadership.",
        },
        "reporting": {
            "Solid Line To": "HR Director",
            "Matrix/Functional Reporting": "Close partnership with Plant Director and site leadership team for workforce and organizational priorities.",
            "Direct Reports (Number, Job Titles)": "HR specialists, HR administrators, and onboarding support resources as assigned",
            "Dotted-line/Influence Relationships": "Production leadership, HSE, Finance, and external HR/payroll or labor-related partners as relevant.",
        },
        "contacts": "Engages daily with site leadership, supervisors, and employees, and regularly with Group HR, Finance, and support providers. Communication ranges from routine people coordination to sensitive employee relations, staffing constraints, and local change management discussions.",
        "innovation_examples": [
            "Improved onboarding and early employee support for plant roles by aligning HR routines more closely with operational realities and line-manager expectations.",
            "Strengthened local HR visibility into absenteeism, workforce pressure points, and employee relations themes to support earlier intervention.",
            "Standardized key plant HR processes while preserving the flexibility needed for shift-based and production-driven workforce demands.",
            "Built stronger manager capability at site level through practical coaching on feedback, issue handling, and people process discipline.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Human Resources, Business Administration, Psychology, or related field preferred.",
            "Work Experience": "5+ years in HR generalist or HR management roles, ideally in manufacturing, logistics, or another operational environment with shift-based workforces.",
            "Languages": "English and Slovak required.",
        },
        "technical": {
            "Technical Skills": [
                "Employee Relations - Advanced",
                "Site Workforce Planning - Advanced",
                "Onboarding & Employee Lifecycle Management - Advanced",
                "Local Labor Law & Compliance - Advanced",
                "Manager Coaching - Advanced",
                "HR Process Execution - Advanced",
                "HR Reporting & Coordination - Intermediate to Advanced",
            ],
            "Business Knowledge Required": "Good understanding of plant operations, shift-based workforce dynamics, and the people implications of production pressure, staffing gaps, and local leadership capability.",
        },
        "competencies": [
            ("Practical Leadership", "Supports managers and employees with a grounded, operationally aware HR approach that works in a live plant setting."),
            ("Problem-Solving", "Navigates employee issues, staffing constraints, and compliance needs with balanced judgment and practicality."),
            ("Communication", "Communicates clearly and credibly with employees, supervisors, and senior local leaders on both routine and sensitive people topics."),
            ("Adaptability", "Responds well to changing workforce demands, shift realities, and urgent local issues without losing process discipline."),
            ("Collaboration", "Builds trust with plant leadership while staying aligned with Group HR standards and broader people priorities."),
            ("Resilience", "Handles pressure, ambiguity, and emotionally difficult cases with steadiness and professionalism."),
        ],
        "problem_examples": [
            "Balance urgent plant hiring needs with candidate quality, onboarding capacity, and labor market constraints.",
            "Support line leaders through difficult employee relations or attendance issues while maintaining policy consistency and legal discipline.",
            "Adapt Group HR processes so they work effectively in a plant environment without creating unnecessary friction or policy drift.",
            "Handle organizational changes at site level where business urgency, employee sentiment, and legal requirements all interact.",
        ],
        "environment": {
            "Work Conditions": "~60-70% office/meeting environment; ~30-40% presence in production and plant support areas.",
            "Travel": "Limited; primarily site-based with periodic Group HR coordination meetings.",
            "Risk": "Normal occupational risk with regular exposure to plant environments (PPE as required). Manages people, compliance, and employee relations risk at site level.",
        },
    },
    {
        "file_name": "HR Manager - Romania Country Organization.docx",
        "identity": {
            "Job Title": "HR Manager",
            "Department/Function": "Human Resources",
            "Location (Site/Region/Group)": "Bucharest, Romania",
            "Reports To (Role)": "HR Director",
            "Direct Reports": "Local HR administrators and support resources as assigned",
            "Role Level": "Country/Local Management",
        },
        "summary": "Leads HR execution for the Romania country organization, supporting commercial, office, and local operational teams with compliant people processes, hiring coordination, employee relations, and organizational support. Translates Group HR standards into practical local execution in a market-facing environment.",
        "responsibilities": [
            ("Country HR Operations", "Own day-to-day HR delivery for the Romania organization, including employee lifecycle coordination, documentation, and local process execution."),
            ("Hiring & Workforce Support", "Coordinate hiring and workforce planning across local business roles in partnership with country leadership and Group HR."),
            ("Employee Relations", "Handle employee relations matters, provide HR guidance to local managers, and support a healthy and compliant working environment."),
            ("Policy Implementation", "Apply Group HR policies within local legal and market realities, ensuring consistency where possible and appropriate local adaptation where necessary."),
            ("Organizational Support", "Support local restructuring, role changes, and capability needs as the country organization evolves."),
            ("Manager Enablement", "Act as the main HR partner to local business leaders on performance, employee handling, communication, and organizational questions."),
            ("Compliance & Administration", "Ensure sound execution of employment documentation, local labor compliance, and HR administration processes."),
            ("Illustrative KPIs", "time-to-fill, onboarding completion, local compliance quality, employee issue resolution, and HR process reliability."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☑ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly affects people operations, leadership support, compliance, and employee experience across the Romania organization.",
            "Decision Autonomy": "Makes local HR operational decisions within approved HR standards and local legal constraints; escalates significant legal, structural, or higher-risk workforce decisions to HR Director.",
        },
        "reporting": {
            "Solid Line To": "HR Director",
            "Matrix/Functional Reporting": "Close partnership with Country Manager Romania and local business leadership.",
            "Direct Reports (Number, Job Titles)": "Local HR administrators and support resources as assigned",
            "Dotted-line/Influence Relationships": "Finance, local operations leadership, external payroll or labor-related advisors, and local managers across functions.",
        },
        "contacts": "Engages frequently with country leadership, managers, employees, and Group HR, and as needed with external administrative or legal partners. Communication ranges from routine employee lifecycle support to more sensitive matters involving local leadership, change, and compliance.",
        "innovation_examples": [
            "Adapted Group HR processes to local market realities while preserving consistency in key governance and employee lifecycle standards.",
            "Improved hiring coordination and onboarding reliability for a lean local organization with shifting business priorities.",
            "Strengthened local people support for managers who needed more structure in employee handling and organizational communication.",
            "Built better visibility into local HR topics for Group HR without overburdening the country organization with unnecessary reporting complexity.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Human Resources, Business Administration, Psychology, or related field preferred.",
            "Work Experience": "5+ years in HR generalist or HR manager roles, ideally in a country organization, sales office, or mixed business environment requiring strong hands-on execution.",
            "Languages": "English and Romanian required.",
        },
        "technical": {
            "Technical Skills": [
                "Country HR Operations - Advanced",
                "Employee Relations - Advanced",
                "Recruitment & Workforce Coordination - Advanced",
                "Local Labor Compliance - Advanced",
                "Manager Support - Advanced",
                "Policy Translation into Local Practice - Advanced",
                "HR Administration & Documentation - Advanced",
            ],
            "Business Knowledge Required": "Good understanding of local market-facing business environments, leadership support needs, and how to apply Group HR standards in a smaller country context.",
        },
        "competencies": [
            ("Ownership", "Takes end-to-end responsibility for local HR execution and follows through reliably in a lean organizational environment."),
            ("Problem-Solving", "Navigates local HR and compliance issues pragmatically while preserving alignment with broader Group expectations."),
            ("Communication", "Builds trust with managers and employees through clear, practical, and locally credible communication."),
            ("Adaptability", "Handles shifting local priorities and broad HR scope without losing consistency in core process execution."),
            ("Collaboration", "Works smoothly between local leadership and Group HR, translating expectations both ways."),
            ("Discretion", "Handles sensitive people matters with professionalism and appropriate confidentiality."),
        ],
        "problem_examples": [
            "Balance Group HR consistency with local legal, cultural, and organizational realities in a smaller country setup.",
            "Support local managers who need broad HR help across hiring, employee issues, and organizational clarity without creating dependency or confusion.",
            "Keep HR administration and compliance strong even when the local organization is lean and priorities shift quickly.",
            "Manage people topics in a business environment where roles are broad and HR often has to solve both strategic and administrative issues at once.",
        ],
        "environment": {
            "Work Conditions": "~80-90% office/meeting environment; limited operational site presence depending on local business needs.",
            "Travel": "Occasional within Romania and periodic coordination with regional or Group HR stakeholders.",
            "Risk": "Normal occupational risk. Manages local people, compliance, and leadership support risks for the Romania organization.",
        },
    },
    {
        "file_name": "Marketing Manager - Vienna HQ.docx",
        "identity": {
            "Job Title": "Marketing Manager",
            "Department/Function": "Marketing / Brand Management",
            "Location (Site/Region/Group)": "Wien, Austria",
            "Reports To (Role)": "Marketing Director",
            "Direct Reports": "No formal direct reports; may coordinate agencies, junior marketers, and cross-functional launch teams",
            "Role Level": "Manager / HQ Functional Role",
        },
        "summary": "Supports the central marketing agenda from Vienna HQ, translating brand and portfolio strategy into coordinated campaigns, launch support, and communication materials across markets. Acts as a cross-functional coordinator between Marketing, Commercial, R&D, and external partners.",
        "responsibilities": [
            ("Brand & Campaign Coordination", "Support the planning and execution of central marketing campaigns, product communication, and brand initiatives aligned with Group priorities."),
            ("Portfolio & Launch Support", "Coordinate launch assets, messaging, and readiness activities in partnership with Commercial, R&D, Quality, and external agencies."),
            ("Content & Communication Development", "Develop or coordinate communication materials, campaign briefs, and brand-supporting content across channels and stakeholders."),
            ("Agency & Partner Management", "Work with creative, media, and production partners to deliver materials on time and to expected quality standards."),
            ("Cross-market Alignment", "Help ensure that central marketing direction is translated into usable guidance for local market teams without creating unnecessary complexity."),
            ("Performance Tracking", "Monitor campaign effectiveness, execution quality, and feedback from local teams to improve future marketing output."),
            ("Operational Marketing Support", "Maintain marketing planning routines, asset management discipline, and coordination points that keep central activity moving efficiently."),
            ("Illustrative KPIs", "campaign delivery timeliness, launch readiness, asset quality, stakeholder satisfaction, and selected brand or activation performance measures."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☑ Multi-site ☐ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Supports central brand and marketing execution with influence across multiple markets through stronger coordination, content quality, and launch readiness.",
            "Decision Autonomy": "Makes day-to-day decisions on campaign coordination, content development, and execution management within approved brand direction and budget boundaries; escalates major strategic or investment decisions to Marketing Director.",
        },
        "reporting": {
            "Solid Line To": "Marketing Director",
            "Matrix/Functional Reporting": "Works closely with VP Commercial, R&D, Commercial market leads, and external agencies depending on campaigns and launch scope.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Local market marketers, brand managers, key account teams, and agency partners.",
        },
        "contacts": "Engages regularly with central marketing stakeholders, commercial leads, R&D, and agencies, and periodically with local market teams. Communication ranges from project coordination and asset review to campaign planning and launch support alignment.",
        "innovation_examples": [
            "Improved central-to-local campaign handovers by clarifying what local teams needed to activate brand initiatives effectively.",
            "Built more structured launch readiness coordination across product, commercial, and marketing stakeholders.",
            "Reduced rework in marketing assets and communication materials by tightening brief quality and approval routines.",
            "Helped central marketing work more effectively with cross-functional stakeholders by making timelines, ownership, and dependencies more visible.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Marketing, Business, Communications, or related field preferred.",
            "Work Experience": "4-7 years in brand marketing, campaign coordination, or category support, ideally in FMCG or another multi-market consumer business.",
            "Languages": "English required; German preferred.",
        },
        "technical": {
            "Technical Skills": [
                "Campaign Coordination - Advanced",
                "Brand Support & Launch Readiness - Advanced",
                "Agency Management - Advanced",
                "Cross-functional Marketing Execution - Advanced",
                "Content Briefing & Asset Coordination - Advanced",
                "Performance Tracking - Intermediate to Advanced",
                "Marketing Planning Discipline - Advanced",
            ],
            "Business Knowledge Required": "Understanding of how central brand direction, local activation needs, product readiness, and commercial timing interact in a multi-market consumer business.",
        },
        "competencies": [
            ("Coordination", "Keeps multiple stakeholders and deliverables moving without losing clarity on dependencies or priorities."),
            ("Problem-Solving", "Finds workable solutions when timelines, brand standards, and local activation needs do not align perfectly."),
            ("Communication", "Translates marketing intent into clear briefs, expectations, and practical action for partners and stakeholders."),
            ("Adaptability", "Handles shifting priorities, launch changes, and fast campaign timelines without losing execution quality."),
            ("Collaboration", "Works comfortably across creative, commercial, product, and operational interfaces."),
            ("Attention to Detail", "Protects brand quality through disciplined review and asset coordination."),
        ],
        "problem_examples": [
            "Support launches where product timing, packaging readiness, and campaign timing are not perfectly aligned.",
            "Balance central brand consistency with practical needs from different local markets.",
            "Keep agencies and internal stakeholders aligned when approvals or inputs arrive late.",
            "Translate broad strategic direction into usable campaign plans and assets with clear ownership.",
        ],
        "environment": {
            "Work Conditions": "~90% office/meeting environment; occasional travel for workshops, shoots, or market coordination.",
            "Travel": "Periodic within Europe depending on launches, workshops, and agency or market meetings.",
            "Risk": "Normal occupational risk. Accountable for execution quality and coordination effectiveness in central marketing activity.",
        },
    },
    {
        "file_name": "Marketing Manager - Romania Local Market.docx",
        "identity": {
            "Job Title": "Marketing Manager",
            "Department/Function": "Marketing / Local Market Activation",
            "Location (Site/Region/Group)": "Bucharest, Romania",
            "Reports To (Role)": "Country Manager Romania",
            "Direct Reports": "No formal direct reports; may coordinate agencies and local activation partners",
            "Role Level": "Manager / Local Market Role",
        },
        "summary": "Leads local marketing execution in Romania, adapting central brand direction to the realities of the local market, customers, and channels. Focuses on activation, campaign delivery, market relevance, and commercial support rather than central portfolio ownership.",
        "responsibilities": [
            ("Local Market Activation", "Plan and execute local campaigns, promotions, and communication activities that support commercial growth and brand presence in Romania."),
            ("Brand Adaptation", "Translate central brand and campaign direction into market-relevant activation suited to local consumers, retail realities, and channel needs."),
            ("Commercial Support", "Partner with local commercial teams to support sell-in, visibility, promotions, trade priorities, and customer-facing initiatives."),
            ("Agency & Partner Coordination", "Coordinate local agencies, media, production, and activation partners to deliver campaigns and in-market execution effectively."),
            ("Market Insight Use", "Incorporate local market feedback, competitor signals, and consumer context into campaign planning and marketing recommendations."),
            ("Budget & Execution Discipline", "Manage local marketing spend and execution plans with strong attention to practical impact, timing, and resource constraints."),
            ("Cross-functional Coordination", "Stay aligned with HQ Marketing, Commercial, and external partners so local execution remains connected to the broader brand direction."),
            ("Illustrative KPIs", "campaign execution quality, activation timeliness, in-market visibility, promotion support effectiveness, and selected local growth indicators."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☑ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly affects local market visibility, execution quality, and the ability to translate brand strategy into commercial traction in Romania.",
            "Decision Autonomy": "Makes local campaign and activation decisions within approved budget and brand boundaries; escalates major strategic shifts or cross-market implications to Country Manager and HQ Marketing.",
        },
        "reporting": {
            "Solid Line To": "Country Manager Romania",
            "Matrix/Functional Reporting": "Dotted-line alignment with Marketing Director for brand consistency and central campaign coordination.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Local commercial teams, agencies, distributors, and HQ marketing stakeholders.",
        },
        "contacts": "Engages frequently with country commercial leadership, HQ marketing, local agencies, and customer-facing teams. Communication ranges from campaign execution and trade support to feedback on what is or is not working in the local market.",
        "innovation_examples": [
            "Improved local campaign effectiveness by adapting central materials more intelligently to local channels and customer needs.",
            "Built stronger coordination between marketing and local commercial teams so activation work better supported sell-in and market visibility.",
            "Introduced more disciplined local planning and agency coordination to reduce last-minute execution issues.",
            "Provided clearer local market feedback into central marketing conversations, improving relevance of future campaigns and materials.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Marketing, Business, Communications, or related field preferred.",
            "Work Experience": "4-7 years in market-facing marketing, activation, brand support, or trade marketing, ideally in FMCG or consumer products.",
            "Languages": "English and Romanian required.",
        },
        "technical": {
            "Technical Skills": [
                "Local Campaign Execution - Advanced",
                "Brand Adaptation for Market Use - Advanced",
                "Commercial & Trade Marketing Support - Advanced",
                "Agency Coordination - Advanced",
                "Budget & Activation Management - Advanced",
                "Market Insight Application - Intermediate to Advanced",
                "Cross-functional Local Execution - Advanced",
            ],
            "Business Knowledge Required": "Good understanding of local market dynamics, customer realities, channel activation, and how central brand direction must be adapted to work in-country.",
        },
        "competencies": [
            ("Ownership", "Takes practical responsibility for making campaigns and activation work in the real local market context."),
            ("Problem-Solving", "Adjusts plans quickly when local market conditions, customer expectations, or resource constraints change."),
            ("Communication", "Bridges HQ brand language and local commercial needs in a clear and constructive way."),
            ("Adaptability", "Handles a fast-moving local environment with shifting priorities and imperfect inputs."),
            ("Collaboration", "Works well across commercial, marketing, and partner relationships where influence matters more than hierarchy."),
            ("Execution Focus", "Keeps marketing grounded in visible, timely, and commercially useful delivery."),
        ],
        "problem_examples": [
            "Adapt central campaigns that do not fully fit local channel or customer realities without weakening the brand too far.",
            "Support commercial goals when budget, timing, and market conditions limit the ideal marketing plan.",
            "Coordinate agencies and internal stakeholders effectively in a lean local organization with broad role scope.",
            "Balance local autonomy with the need to stay aligned to Group brand and campaign direction.",
        ],
        "environment": {
            "Work Conditions": "~85-90% office/meeting environment; periodic in-market visits, retail observation, and partner meetings.",
            "Travel": "Occasional within Romania and periodic coordination with HQ or regional stakeholders.",
            "Risk": "Normal occupational risk. Accountable for local campaign execution quality and market activation effectiveness.",
        },
    },
    {
        "file_name": "Key Account Manager - Czech Republic.docx",
        "identity": {
            "Job Title": "Key Account Manager",
            "Department/Function": "Commercial / Sales",
            "Location (Site/Region/Group)": "Brno, Czech Republic",
            "Reports To (Role)": "Sales Director for Brand Products",
            "Direct Reports": "No direct reports",
            "Role Level": "Manager / Local Market Commercial Role",
        },
        "summary": "Owns key customer relationships in the Czech market, balancing commercial growth, customer negotiation, forecasting input, and execution discipline. Acts as the main commercial interface for assigned accounts and translates customer priorities into actionable internal coordination.",
        "responsibilities": [
            ("Account Ownership", "Manage the end-to-end relationship with assigned key accounts in the Czech Republic, including customer planning, negotiation, and day-to-day commercial coordination."),
            ("Joint Business Planning", "Develop account plans, growth opportunities, and promotional priorities aligned with company objectives and customer potential."),
            ("Negotiation & Commercial Delivery", "Lead customer negotiations on pricing, assortment, promotions, visibility, and trading terms within approved boundaries."),
            ("Forecast & Demand Input", "Provide grounded input into demand expectations, customer changes, and promotional impact to support supply and planning alignment."),
            ("Cross-functional Coordination", "Work closely with marketing, supply chain, finance, and customer service to ensure customer commitments are realistic and well supported."),
            ("Execution Monitoring", "Track account performance, customer compliance, and commercial execution quality; identify risks and opportunities early."),
            ("Relationship Development", "Build trusted relationships with customer stakeholders and represent the company with professionalism and commercial discipline."),
            ("Illustrative KPIs", "sales growth, account profitability, promotional effectiveness, forecast quality, customer satisfaction, and execution against account plans."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☑ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly influences revenue, customer relationships, market visibility, and execution quality in the Czech market.",
            "Decision Autonomy": "Makes day-to-day customer and negotiation decisions within approved commercial strategy and financial boundaries; escalates major contractual or strategic account issues to sales leadership.",
        },
        "reporting": {
            "Solid Line To": "Sales Director for Brand Products",
            "Matrix/Functional Reporting": "Works closely with Marketing, Supply Chain, Finance, and Customer Service on customer plans and delivery risks.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Trade marketing, planning, customer service, and distributor or retailer stakeholders as applicable.",
        },
        "contacts": "Engages frequently with customer buyers, internal commercial leadership, customer service, and cross-functional support teams. Communication ranges from negotiation and customer planning to issue resolution and forecast-related coordination.",
        "innovation_examples": [
            "Built stronger customer planning routines that improved account visibility and reduced surprises in promotional and supply execution.",
            "Improved cross-functional communication around customer activity so internal teams could support account commitments more effectively.",
            "Identified pragmatic assortment or activation opportunities that strengthened customer relationships without adding unnecessary complexity.",
            "Created more structured follow-up on account performance and customer actions to support better commercial prioritization.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Business, Sales, Marketing, Economics, or related field preferred.",
            "Work Experience": "4-7 years in sales, key account management, or commercial roles, ideally in FMCG or another customer-driven market environment.",
            "Languages": "English and Czech required.",
        },
        "technical": {
            "Technical Skills": [
                "Key Account Management - Advanced",
                "Negotiation & Trade Terms Management - Advanced",
                "Joint Business Planning - Advanced",
                "Forecast & Demand Input - Advanced",
                "Cross-functional Commercial Coordination - Advanced",
                "Commercial Performance Tracking - Advanced",
                "Customer Relationship Development - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of customer dynamics, retail expectations, promotional mechanics, and the operational implications of commercial commitments in the Czech market.",
        },
        "competencies": [
            ("Commercial Ownership", "Takes strong responsibility for customer outcomes and follows through on both opportunity and risk."),
            ("Problem-Solving", "Balances customer expectations with internal operational and financial realities in a practical way."),
            ("Communication", "Communicates clearly with customers and internal teams, especially when trade-offs or issues need to be managed quickly."),
            ("Adaptability", "Responds well to shifting customer priorities, market pressure, and forecast or service volatility."),
            ("Collaboration", "Works effectively across internal functions where customer success depends on coordinated follow-through."),
            ("Resilience", "Handles negotiation pressure, performance swings, and difficult customer situations with steady judgment."),
        ],
        "problem_examples": [
            "Manage customer requests that create tension between commercial growth goals and supply or margin realities.",
            "Keep forecast input and promotion plans useful even when customer behavior is volatile or imperfectly communicated.",
            "Negotiate effectively without overcommitting the organization operationally or financially.",
            "Maintain strong customer trust while navigating service issues, delays, or commercial disagreements.",
        ],
        "environment": {
            "Work Conditions": "~80-90% office/meeting/customer environment; periodic travel to customers and trade-related meetings.",
            "Travel": "Frequent within the Czech Republic for customer meetings and occasional regional coordination.",
            "Risk": "Normal occupational risk. Accountable for commercial execution, customer relationship quality, and account performance in the local market.",
        },
    },
    {
        "file_name": "Key Account Manager - Poland Modern Trade.docx",
        "identity": {
            "Job Title": "Key Account Manager",
            "Department/Function": "Commercial / Sales",
            "Location (Site/Region/Group)": "Warszawa, Poland",
            "Reports To (Role)": "Commercial Director Poland",
            "Direct Reports": "No direct reports",
            "Role Level": "Manager / Channel-Focused Commercial Role",
        },
        "summary": "Owns modern trade key accounts in Poland, focusing on structured customer development, commercial negotiation, promotional planning, and channel execution. Operates in a more concentrated retailer environment with strong emphasis on account discipline, promotional mechanics, and execution quality.",
        "responsibilities": [
            ("Modern Trade Account Ownership", "Manage the relationship, growth agenda, and commercial execution for assigned modern trade customers in Poland."),
            ("Channel Planning", "Build account plans and activity calendars that reflect modern trade promotional cycles, customer expectations, and channel economics."),
            ("Negotiation & Terms Management", "Lead negotiations on pricing, assortment, promotions, trade support, and execution conditions within approved commercial boundaries."),
            ("Promotional & Demand Coordination", "Coordinate promotional assumptions and customer plans with supply chain and planning teams to improve readiness and reduce avoidable execution risk."),
            ("Performance Management", "Monitor account performance, customer-specific profitability, execution quality, and activity outcomes to improve commercial decision-making."),
            ("Cross-functional Coordination", "Partner with trade marketing, customer service, supply chain, and finance to support customer commitments and resolve issues effectively."),
            ("Relationship Development", "Build durable working relationships with customer stakeholders and maintain a credible commercial presence in a competitive modern trade environment."),
            ("Illustrative KPIs", "modern trade revenue growth, promotional ROI, account profitability, plan execution quality, service support, and forecast reliability."),
        ],
        "org_scope": {
            "Sphere of Influence": "☐ Enterprise/Group ☐ Multi-site ☑ Regional ☐ Site/Local ☐ Team",
            "Extent of Impact": "Directly affects revenue, visibility, and execution quality within a strategically important modern trade channel in Poland.",
            "Decision Autonomy": "Makes day-to-day account and negotiation decisions within approved commercial strategy and channel guardrails; escalates major customer, financial, or strategic decisions to commercial leadership.",
        },
        "reporting": {
            "Solid Line To": "Commercial Director Poland",
            "Matrix/Functional Reporting": "Works closely with Trade Marketing, Supply Chain, Finance, and Customer Service on channel execution and customer support.",
            "Direct Reports (Number, Job Titles)": "No formal direct reports",
            "Dotted-line/Influence Relationships": "Planning, trade marketing, customer service, and customer-side category or buying stakeholders.",
        },
        "contacts": "Engages frequently with modern trade customers, local commercial leadership, trade marketing, and execution support teams. Communication ranges from negotiation and account planning to promotional coordination and escalation of service or execution issues.",
        "innovation_examples": [
            "Improved promotional planning discipline for modern trade customers by strengthening alignment between customer calendars and supply readiness.",
            "Built clearer customer-specific performance visibility that supported better decision-making on promotions, assortment, and account focus.",
            "Improved collaboration between commercial and support functions to reduce execution gaps during high-activity customer periods.",
            "Developed more structured account planning that reflected the specific needs and economics of modern trade in the Polish market.",
        ],
        "qualifications": {
            "Education": "Bachelor's degree in Business, Sales, Marketing, Economics, or related field preferred.",
            "Work Experience": "4-7 years in key account management, sales, or trade-facing commercial roles, ideally with exposure to modern trade or large retail environments.",
            "Languages": "English and Polish required.",
        },
        "technical": {
            "Technical Skills": [
                "Key Account Management - Advanced",
                "Modern Trade Channel Planning - Advanced",
                "Negotiation & Promotional Planning - Advanced",
                "Account Performance Management - Advanced",
                "Cross-functional Commercial Coordination - Advanced",
                "Forecast & Activity Alignment - Advanced",
                "Customer Relationship Development - Advanced",
            ],
            "Business Knowledge Required": "Strong understanding of modern trade dynamics, customer planning cycles, promotional economics, and the operational implications of retail-facing commercial commitments in Poland.",
        },
        "competencies": [
            ("Commercial Discipline", "Brings strong structure to account planning, negotiation, and follow-through in a demanding retail channel."),
            ("Problem-Solving", "Works through commercial and execution issues with an eye on both customer needs and internal constraints."),
            ("Communication", "Communicates clearly in negotiation, planning, and issue-resolution settings with both customers and internal teams."),
            ("Adaptability", "Handles high activity levels, changing retailer demands, and time-sensitive promotional execution well."),
            ("Collaboration", "Builds effective partnerships across support functions where channel performance depends on coordinated action."),
            ("Persistence", "Maintains follow-through and account focus in a channel where execution details materially affect results."),
        ],
        "problem_examples": [
            "Balance aggressive retail expectations with supply, margin, and internal capability constraints.",
            "Keep promotional activity aligned across customer, commercial, and supply teams in a channel with tight timelines and high execution pressure.",
            "Prioritize the right modern trade opportunities when multiple customers compete for attention and resources at the same time.",
            "Protect long-term account quality while still delivering short-term results in a highly demanding retail environment.",
        ],
        "environment": {
            "Work Conditions": "~80-90% office/meeting/customer environment; periodic travel to retail customers and market meetings.",
            "Travel": "Frequent within Poland for customer meetings and periodic commercial coordination.",
            "Risk": "Normal occupational risk. Accountable for channel execution quality, customer outcomes, and commercial performance in modern trade.",
        },
    },
]


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Aptos"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    heading1.font.size = Pt(24)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.space_before = Pt(16.1)
    heading1.paragraph_format.space_after = Pt(16.1)


def add_field_paragraph(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="Normal")
    r1 = p.add_run(f"{label}:")
    r1.bold = True
    r2 = p.add_run(f" {value}")
    r2.bold = False


def add_labeled_paragraph(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="Normal")
    r1 = p.add_run(f"{label}:")
    r1.bold = True
    r2 = p.add_run(f" {value}")
    r2.bold = False


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def build_doc(job: dict) -> Document:
    doc = Document()
    configure_styles(doc)

    add_heading(doc, "A. Job Identity")
    for label, value in job["identity"].items():
        add_field_paragraph(doc, label, value)

    add_heading(doc, "B. Role Summary (Scope and Impact)")
    doc.add_paragraph(job["summary"], style="Normal")

    add_heading(doc, "C. Key Responsibilities")
    for label, value in job["responsibilities"]:
        add_labeled_paragraph(doc, label, value)

    add_heading(doc, "D. Organizational Scope")
    for label, value in job["org_scope"].items():
        add_field_paragraph(doc, label, value)

    add_heading(doc, "E. Reporting Relationships")
    for label, value in job["reporting"].items():
        add_field_paragraph(doc, label, value)

    add_heading(doc, "F. Key Contacts and Communication")
    doc.add_paragraph(job["contacts"], style="Normal")

    add_heading(doc, "G. Contribution to Innovation and Change - Examples")
    for item in job["innovation_examples"]:
        doc.add_paragraph(item, style="Normal")

    add_heading(doc, "H. Qualifications")
    for label, value in job["qualifications"].items():
        add_field_paragraph(doc, label, value)

    add_heading(doc, "I. Technical & Business Knowledge")
    doc.add_paragraph("Technical Skills:", style="Normal")
    for item in job["technical"]["Technical Skills"]:
        doc.add_paragraph(item, style="Normal")
    add_field_paragraph(doc, "Business Knowledge Required", job["technical"]["Business Knowledge Required"])

    add_heading(doc, "J. Competencies & Personal Attributes")
    for label, value in job["competencies"]:
        add_labeled_paragraph(doc, label, value)

    add_heading(doc, "K. Problem Solving & Decision Complexity - Examples")
    for item in job["problem_examples"]:
        doc.add_paragraph(item, style="Normal")

    add_heading(doc, "L. Environment & Risk")
    for label, value in job["environment"].items():
        add_field_paragraph(doc, label, value)

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for job in JOB_DESCRIPTIONS:
        doc = build_doc(job)
        out_path = OUTPUT_DIR / job["file_name"]
        doc.save(out_path)
        print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
